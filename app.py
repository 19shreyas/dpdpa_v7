import streamlit as st
import openai
import json
import pandas as pd
import re
import fitz
import io
import datetime
from docx import Document              

# --- OpenAI Setup ---
api_key = st.secrets["OPENAI_API_KEY"]
client = openai.OpenAI(api_key=api_key)

# --- Section Checklists ---
dpdpa_checklists = {
    "4": {
        "title": "Grounds for Processing Personal Data",
        "items": [
            {"id" : "4.1", "text" : "The policy must state that personal data is processed **only as per the provisions of the Digital Personal Data Protection Act, 2023**."},
            {"id" : "4.2", "text" : "The policy must confirm that personal data is processed **only for a lawful purpose**."},
            {"id" : "4.3", "text" : "The policy must define **lawful purpose** as any purpose **not expressly forbidden by law**."},
            {"id" : "4.4", "text" : "The policy must include a statement that personal data is processed **only with the consent of the Data Principal**."},
            {"id" : "4.5", "text" : "Alternatively, the policy must specify that personal data is processed **for certain legitimate uses**, as defined under the Act."}
        ]
    },
    "5": {
        "title": "Notice",
        "items": [
            {"id" : "5.1", "text" : "The policy must state that **every request for consent** is accompanied or preceded by a **notice from the Data Fiduciary to the Data Principal**."},
            {"id" : "5.2", "text" : "The notice must clearly specify the **personal data proposed to be processed**."},
            {"id" : "5.3", "text" : "The notice must clearly specify the **purpose for which the personal data is proposed to be processed**."},
            {"id" : "5.4", "text" : "The notice must explain the **manner in which the Data Principal can exercise her rights under Section 6(4)** (withdrawal of consent)."},
            {"id" : "5.5", "text" : "The notice must explain the **manner in which the Data Principal can exercise her rights under Section 13** (grievance redressal)."},
            {"id" : "5.6", "text" : "The notice must specify the **manner in which a complaint can be made to the Data Protection Board**."},
            {"id" : "5.7", "text" : "If consent was obtained **before the commencement of the Act**, the policy must state that a notice will be sent **as soon as reasonably practicable**."},
            {"id" : "5.8", "text" : "The post-commencement notice must mention the **personal data that has been processed**."},
            {"id" : "5.9", "text" : "The post-commencement notice must mention the **purpose for which the personal data has been processed**."},
            {"id" : "5.10", "text" : "The post-commencement notice must mention the **manner in which the Data Principal can exercise her rights under Section 6(4)**."},
            {"id" : "5.11", "text" : "The post-commencement notice must mention the **manner in which the Data Principal can exercise her rights under Section 13**."},
            {"id" : "5.12", "text" : "The post-commencement notice must mention the **manner in which a complaint can be made to the Board**."},
            {"id" : "5.13", "text" : "The policy must mention that the Data Fiduciary **may continue to process personal data** until the Data Principal **withdraws her consent**."},
            {"id" : "5.14", "text" : "The policy must provide the Data Principal an **option to access the contents of the notice** in **English or any language listed in the Eighth Schedule of the Constitution**."}
        ]
    },
    "6": {
        "title": "Consent",
        "items": [
            {"id" : "6.1", "text" : "The policy must state that **consent is free, specific, informed, unconditional, and unambiguous**, given through a **clear affirmative action**."},
            {"id" : "6.2", "text" : "The policy must specify that **consent signifies agreement to process personal data only for the specified purpose**."},
            {"id" : "6.3", "text" : "The policy must state that **consent is limited to such personal data as is necessary for the specified purpose**."},
            {"id" : "6.4", "text" : "The policy must mention that **any part of the consent that violates this Act, rules under it, or any other law in force is invalid to that extent**."},
            {"id" : "6.5", "text" : "The request for consent must be presented in **clear and plain language**."},
            {"id" : "6.6", "text" : "The request for consent must allow the Data Principal to access it in **English or any language listed in the Eighth Schedule of the Constitution**."},
            {"id" : "6.7", "text" : "The request for consent must provide **contact details of a Data Protection Officer** or **another authorised person** responsible for handling Data Principal queries."},
            {"id" : "6.8", "text" : "The policy must clearly state that the **Data Principal has the right to withdraw consent at any time**."},
            {"id" : "6.9", "text" : "The **ease of withdrawing consent** must be comparable to the **ease with which consent was given**."},
            {"id" : "6.10", "text" : "The policy must mention that **consequences of withdrawal shall be borne by the Data Principal**."},
            {"id" : "6.11", "text" : "The policy must state that **withdrawal does not affect the legality of data processing done before withdrawal**."},
            {"id" : "6.12", "text" : "The policy must mention that upon withdrawal of consent, the **Data Fiduciary and its Data Processors must cease processing** the personal data **within a reasonable time**, unless permitted by law."},
            {"id" : "6.13", "text" : "The policy must state that consent **can be managed, reviewed, or withdrawn through a Consent Manager**."},
            {"id" : "6.14", "text" : "The policy must specify that the **Consent Manager is accountable to the Data Principal** and acts on her behalf."},
            {"id" : "6.15", "text" : "The policy must specify that **every Consent Manager is registered with the Board** under prescribed conditions."},
            {"id" : "6.16", "text" : "The policy must mention that, in case of dispute, the **Data Fiduciary must prove that proper notice was given and valid consent was obtained** as per the Act and its rules."}
        ]
    },
    "7": {
        "title": "Certain Legitimate Uses",
        "items": [
            {"id" : "7.1", "text" : "The policy must allow personal data to be processed for the **specified purpose for which the Data Principal voluntarily provided the data**, if she has **not indicated non-consent** to such use."},
            {"id" : "7.2", "text" : "The policy must permit personal data to be processed by the State or its instrumentalities for providing or issuing **subsidy, benefit, service, certificate, licence, or permit**, as prescribed, where the Data Principal has **previously consented** to such processing."},
            {"id" : "7.3", "text" : "The policy must allow personal data to be processed by the State or its instrumentalities if the data is **already available in digital or digitised form in notified government databases**, subject to prescribed standards and government policies."},
            {"id" : "7.4", "text" : "The policy must allow personal data to be processed by the State or its instrumentalities for performing any **legal function** under existing Indian laws or **in the interest of sovereignty and integrity of India or State security**."},
            {"id" : "7.5", "text" : "The policy must allow personal data to be processed to **fulfil a legal obligation** requiring any person to disclose information to the State or its instrumentalities, as per applicable laws."},
            {"id" : "7.6", "text" : "The policy must permit personal data to be processed for **compliance with any judgment, decree, or order** issued under Indian law, or for **contractual or civil claims under foreign laws**."},
            {"id" : "7.7", "text" : "The policy must allow personal data to be processed to **respond to a medical emergency** involving a **threat to life or immediate health risk** of the Data Principal or any individual."},
            {"id" : "7.8", "text" : "The policy must allow personal data to be processed to **provide medical treatment or health services** during an **epidemic, outbreak, or other threat to public health**."},
            {"id" : "7.9", "text" : "The policy must permit processing of personal data to **ensure safety of or provide assistance/services to individuals** during any **disaster or breakdown of public order**."},
            {"id" : "7.10", "text" : "The policy must define 'disaster' in accordance with the **Disaster Management Act, 2005 (Section 2(d))**."},
            {"id" : "7.11", "text" : "The policy must allow personal data to be processed for purposes related to **employment**, or to **safeguard the employer from loss or liability**, including prevention of corporate espionage, confidentiality of trade secrets or IP, and enabling services/benefits to employee Data Principals."}
        ]
    },
    "8": {
        "title": "General Obligations of Data Fiduciary",
        "items": [
            {"id" : "8.1", "text" : "The policy must state that the Data Fiduciary is responsible for complying with the Act and its rules, even if the Data Principal fails to perform her duties."},
            {"id" : "8.2", "text" : "The policy must state that the Data Fiduciary may engage or involve a Data Processor **only under a valid contract** to process personal data for offering goods or services."},
            {"id" : "8.3", "text" : "The policy must ensure that if personal data is used to make a decision affecting the Data Principal, the data must be **complete, accurate, and consistent**."},
            {"id" : "8.4", "text" : "The policy must ensure that if personal data is disclosed to another Data Fiduciary, the data must be **complete, accurate, and consistent**."},
            {"id" : "8.5", "text" : "The policy must require the Data Fiduciary to implement **appropriate technical and organisational measures** to ensure compliance with the Act and its rules."},
            {"id" : "8.6", "text" : "The policy must mandate **reasonable security safeguards** to protect personal data from breaches, including breaches by its Data Processors."},
            {"id" : "8.7", "text" : "The policy must state that in the event of a **personal data breach**, the Data Fiduciary shall **inform both the Board and each affected Data Principal** in the prescribed manner."},
            {"id" : "8.8", "text" : "The policy must mandate that personal data be **erased upon withdrawal of consent** or as soon as it is reasonable to assume that the **specified purpose is no longer being served**, whichever is earlier."},
            {"id" : "8.9", "text" : "The policy must mandate that the Data Fiduciary must **cause its Data Processors to erase the data** when retention is no longer justified."},
            {"id" : "8.10", "text" : "The policy must define that the specified purpose is deemed no longer served if the Data Principal has neither **approached the Data Fiduciary for the purpose** nor **exercised her rights** within the prescribed time period."},
            {"id" : "8.11", "text" : "The policy must require publishing the **business contact details** of the Data Protection Officer (if applicable) or of an authorised person able to respond to questions about personal data processing."},
            {"id" : "8.12", "text" : "The policy must provide an **effective grievance redressal mechanism** for Data Principals."},
            {"id" : "8.13", "text" : "The policy must clarify that a Data Principal is considered as **not having approached** the Data Fiduciary if she has not initiated contact in person, or through physical or electronic communication, for the purpose within a prescribed period."}
        ]
    }
}

# --- PDF Extractor ---
def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    return "\n".join(page.get_text() for page in doc)

# --- Prompt Generator ---
def create_full_policy_prompt(section_id, full_policy_text, checklist):
    checklist_text = "\n".join(
        f"{item['id']}. {item['text']}" for item in checklist
    )

    return f"""
    You are a compliance analyst evaluating whether the following full privacy policy meets DPDPA Section {section_id}: {dpdpa_checklists[section_id]['title']}.
    
    **Checklist:** Use the item numbers (e.g., 4.1, 4.2...) from the checklist below in your response. Do not rephrase or modify the checklist items. Evaluate strictly based on the original items.
    
    {checklist_text}
    
    **Full Policy Text:**
    {full_policy_text}
    
    Instructions:
    For each checklist item, search anywhere in the policy and classify it as:
    - Explicitly Mentioned
    - Partially Mentioned
    - Missing
    
    Return output in this JSON format only:
    {{
      "Checklist Evaluation": [
        {{
          "Checklist Item ID": "4.1",
          "Status": "Explicitly Mentioned",
          "Justification": "..."
        }},
        ...
      ],
      "Match Level": "Fully Compliant / Partially Compliant / Non-Compliant",
      "Compliance Score": 0.0,
      "Suggested Rewrite": "...",
      "Simplified Legal Meaning": "..."
    }}
    
    Only return the JSON object. Do not include any commentary or explanation.
    """
# --- GPT Call ---
def call_gpt(prompt, model="gpt-4"):
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    return json.loads(response.choices[0].message.content)
    
def call_gpt_text(prompt, model="gpt-4"):
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

def analyze_policy_section(section_id, checklist, policy_text, model="gpt-4"):
    prompt = create_full_policy_prompt(section_id, policy_text, checklist)
    
    try:
        result = call_gpt(prompt, model=model)
    except Exception as e:
        return {
            "Section": section_id,
            "Title": dpdpa_checklists[section_id]['title'],
            "Error": str(e),
            "Match Level": "Error",
            "Compliance Score": 0.0,
            "Matched Details": [],
            "Checklist Items Matched": [],
            "Suggested Rewrite": "",
            "Simplified Legal Meaning": ""
        }

    checklist_dict = {item["id"]: item["text"] for item in checklist}
    evaluations = []

    matched_count = 0
    partial_count = 0

    for item in result.get("Checklist Evaluation", []):
        item_id = item.get("Checklist Item ID", "").strip()
        status = item.get("Status", "Missing").strip()
        justification = item.get("Justification", "").strip()
        text = checklist_dict.get(item_id, "❓")

        if status == "Explicitly Mentioned":
            matched_count += 1
        elif status == "Partially Mentioned":
            partial_count += 1

        evaluations.append({
            "Checklist Item ID": item_id,
            "Checklist Text": text,
            "Status": status,
            "Justification": justification
        })

    score = (matched_count + 0.5 * partial_count) / len(checklist) if checklist else 0
    level = (
        "Fully Compliant" if score == 1 else
        "Non-Compliant" if score == 0 else
        "Partially Compliant"
    )

    return {
        "Section": section_id,
        "Title": dpdpa_checklists[section_id]['title'],
        "Match Level": result.get("Match Level", level),
        "Compliance Score": round(score, 2),
        "Matched Details": evaluations,
        "Checklist Items Matched": [f"{e['Checklist Item ID']} — {e['Checklist Text']}" for e in evaluations if e["Status"] in ["Explicitly Mentioned", "Partially Mentioned"]],
        "Suggested Rewrite": result.get("Suggested Rewrite", ""),
        "Simplified Legal Meaning": result.get("Simplified Legal Meaning", "")
    }

def set_custom_css():
    st.markdown("""
    <style>
    /* Sidebar Styling */
    section[data-testid="stSidebar"] {
        background-color: #2E2E38;
        color: white !important;
    }
    section[data-testid="stSidebar"] * {
        color: white !important;
    }

    /* Main text and headers */
    section.main div.block-container {
        color: #2E2E38 !important;
    }
    h1, h2, h3, h4, h5, h6, p, span, label, div {
        color: #2E2E38 !important;
    }

    /* Fix for input text color */
    input, textarea, select, div[role="textbox"] {
        color: #2E2E38 !important;
        caret-color: #2E2E38
    }

    /* Fix for selectbox and dropdowns */
    .stSelectbox div, .stRadio div, .stCheckbox div, .stTextInput div, .stDownloadButton div {
        color: #2E2E38 !important;
    }

    /* Fix file uploader text inside black box */
    .stFileUploader div, .stFileUploader span {
    color: white !important;
    }
    /* --- Force white text in selectbox/multiselect dropdowns with dark bg --- */
    div[data-baseweb="select"] {
        color: white !important;
    }
    
    div[data-baseweb="select"] * {
        color: white !important;
    }

    /* Fix Browse files button text inside file uploader */
    .stFileUploader button,
    .stFileUploader label {
    color: black !important;
    font-weight: 500;
    }

    /* Fix download and regular buttons */
    .stButton > button, .stDownloadButton > button {
        color: white !important;
        background-color: #1a9afa;
        font-weight: 600;
        border-radius: 6px;
        border: none;
    }
        
    /* Only for actual input and textarea fields */
    input, textarea {
        background-color: #FFFFFF !important;
        color: #2E2E38 !important;
        border: 1px solid #2E2E38 !important;
        border-radius: 6px !important;
        caret-color: #2E2E38
    }
    /* Additional fix for Streamlit's internal markdown editor (if used) */
    div[contenteditable="true"] {
        caret-color: #2E2E38 !important;
        color: #2E2E38 !important;
    }

    /* For Streamlit selectboxes */
    div[data-baseweb="select"] {
        background-color: #2E2E38 !important;
        color: white !important;
        border-radius: 6px !important;
        border: none !important;
    }
    div[data-baseweb="select"] * {
        color: white !important;
    }

    /* === Fix Streamlit checkbox visibility === */
    div[data-baseweb="checkbox"] > label > div {
        background-color: white !important;
        border: 2px solid #2E2E38 !important;
        border-radius: 4px !important;
        width: 20px !important;
        height: 20px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        margin-right: 10px !important;
    }
    
    /* Show checkmark when selected */
    div[data-baseweb="checkbox"] svg {
        stroke: #2E2E38 !important;
        stroke-width: 2.5px !important;
    }
    /* === FIX: Cursor visibility in text areas === */
    textarea,
    .stTextArea textarea,
    div[role="textbox"],
    .stTextInput input {
        background-color: #FFFFFF !important;
        color: #2E2E38 !important;
        caret-color: #2E2E38 !important;
        border: 1px solid #2E2E38 !important;
        border-radius: 6px !important;
    }
    
    /* Explicit caret fix for editable areas */
    div[contenteditable="true"] {
        caret-color: #2E2E38 !important;
        color: #2E2E38 !important;
    }

    /* Force DataFrame cell background and text to be readable */
.css-1r6slb0 .element-container {
        background-color: white !important;
        color: black !important;
    }

    </style>
    """, unsafe_allow_html=True)

# --- Sidebar Navigation ---
st.set_page_config(page_title="DPDPA Compliance Tool", layout="wide")
set_custom_css()
st.sidebar.markdown("<h1 style='font-size:42px; font-weight:700;'>Navigation</h1>", unsafe_allow_html=True)

menu = st.sidebar.radio("", [
    "Homepage",
    "Policy Compliance Checker",
    "Policy Generator",
    "Dashboard & Reports",
    "Knowledge Assistant",
    "Admin Settings"
])
st.sidebar.markdown("<br><br><br><br><br><br><br><br><br><br><br><br><br><br>", unsafe_allow_html=True)
st.sidebar.markdown("""
    <div style='padding: 0px 12px 0px 0px;'>
        <img src='https://i.postimg.cc/j2dv9kZ2/EY-Parthenon-idp-Wq1a8hl-0.png' width='200'>
    </div>
""", unsafe_allow_html=True)
# --- Homepage ---
if menu == "Homepage":
    st.title("DPDPA Compliance Tool")
    st.markdown("""
    Welcome to the Digital Personal Data Protection Act (DPDPA) Compliance Platform.
    Use the navigation panel to begin generating or matching your policy to India's latest data protection laws.
    """)

# --- Policy Generator ---
elif menu == "Policy Generator":
    st.title("Create a new Policy")
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "Full Policy Generator", "Section-wise Generator", "Lifecycle-wise Template", 
        "GPT Draft Assistant", "Saved Drafts"])

    with tab1:
        st.markdown("### Generate a DPDPA-Compliant Policy")
        st.caption("All fields marked with * are mandatory. This tool creates a draft policy aligned with India's Digital Personal Data Protection Act, 2023.")
    
        # --- Group 1: Organization Details ---
        with st.expander("Organization Details", expanded=False):
            st.markdown("**Policy Type***  \n_Select the kind of policy you are generating._")
            policy_type = st.selectbox(" ", [
                "-- Select Policy Type --", "Privacy Policy", "Retention Policy", "Security Policy", "Data Protection Policy"
            ], key="policy_type")
        
            st.markdown("**Organization Name***  \n_Enter the full legal name of your organization._")
            org_name = st.text_input(" ", key="org_name")
        
            st.markdown("**Sector***  \n_Select your sector or enter a custom one below._")
            sector_dropdown = st.selectbox(" ", [
                "-- Select Sector --", "Healthcare", "Finance", "Education", "Retail", "Government", "Technology", "Telecom"
            ], key="sector_dropdown")
            sector_custom = st.text_input("Custom sector (if not listed)", key="sector_custom")

        # --- Group 2: Data Collection & Target Audience ---
        with st.expander("Data Collection Details", expanded=False):
            st.markdown("**Select Data Types Collected***  \n_Choose the categories of personal data you collect. Required under Sections 6 & 7._")
            data_types_common = st.multiselect(" ", [
                "Name", "Email", "Phone Number", "Biometric", "Health Records", "Location Data",
                "Financial Information", "Browsing Data"
            ], key="data_types_common")
        
            st.markdown("Custom data types (comma separated)")
            data_types_custom = st.text_input(" ", key="data_types_custom")
        
            st.markdown("**Is this policy applicable to children under 18?**  \n_If your services are used by minors, Section 9 applies._")
            children_data = st.radio(" ", ["No", "Yes"], horizontal=True, key="children_data")

        # --- Group 3: Legal Basis & Consent ---
        with st.expander("Legal Basis for Processing", expanded=False):
            st.markdown("**Lawful Purpose for Data Collection***  \n_Why are you collecting personal data? e.g., Delivering services, fraud prevention._")
            lawful_purpose = st.text_input(" ", key="lawful_purpose")
        
            st.markdown("**Type of Consent Taken***  \n_How do you obtain consent? Explicit (opt-in), Deemed (via use), or just Notice._")
            consent_type = st.selectbox(" ", [
                "Explicit Consent", "Deemed Consent", "Notice Only (limited use cases)"
            ], key="consent_type")
        
            st.markdown("**Special Purpose under Section 7 (Optional)**  \n_If you process data without consent for employment, emergency, or public interest, select here._")
            legitimate_use = st.multiselect(" ", [
                "Employment Purposes", "Medical Emergency", "Government Function", "Disaster Response", "Public Interest", "None of the Above"
            ], key="legitimate_use")

        # --- Group 4: Retention, Security, and Sharing ---
        with st.expander("Retention & Data Sharing", expanded=False):
            st.markdown("**Data Retention Duration***  \n_How long do you retain data? Must be justified under Section 8(7)._")
            retention_period = st.text_input(" ", key="retention_period")
        
            st.markdown("**Is Data Shared Internationally?***  \n_Required under Section 16. Mention if data is processed outside India._")
            cross_border = st.radio(" ", ["Yes", "No"], horizontal=True, key="cross_border")
    
        # --- Group 5: Grievance Contact Info ---
        with st.expander("Grievance Redressal Contact", expanded=False):
            st.markdown("**Grievance Officer Contact Email***  \n_As per Section 10, provide a contact for complaints or data requests._")
            grievance_email = st.text_input(" ", key="grievance_email")

        # --- Generate Button ---
        if st.button("Generate DPDPA-Compliant Policy with GPT"):
            errors = []
            if policy_type == "-- Select Policy Type --":
                errors.append("Policy Type")
            if not org_name.strip():
                errors.append("Organization Name")
            if sector_dropdown == "-- Select Sector --" and not sector_custom.strip():
                errors.append("Sector")
            if not data_types_common and not data_types_custom.strip():
                errors.append("Data Types")
            if not lawful_purpose.strip():
                errors.append("Lawful Purpose")
            if not retention_period.strip():
                errors.append("Data Retention Duration")
            if not grievance_email.strip():
                errors.append("Grievance Officer Contact Email")
    
            if errors:
                st.error("Please fill the following mandatory fields: " + ", ".join(errors))
            else:
                sector_final = sector_custom if sector_custom else sector_dropdown
                data_types_final = data_types_common + [dt.strip() for dt in data_types_custom.split(",") if dt.strip()]
                special_uses = ", ".join(legitimate_use) if legitimate_use else "None"
    
                with st.spinner("Generating policy... please wait."):
                    prompt = f"""
    You are a legal policy assistant. Draft a comprehensive, DPDPA-compliant {policy_type.lower()} for the following organization.
    
    **Organization Details**:
    - Name: {org_name}
    - Sector: {sector_final}
    
    **Data Handling**:
    - Data Types Collected: {", ".join(data_types_final)}
    - Applicable to Children (<18): {children_data}
    - Shared Internationally: {cross_border}
    
    **Legal Basis**:
    - Lawful Purpose: {lawful_purpose}
    - Consent Type: {consent_type}
    - Special Use Cases (Sec 7): {special_uses}
    
    **Retention & Redressal**:
    - Retention Period: {retention_period}
    - Grievance Contact: {grievance_email}
    
    **Compliance Requirements**:
    Ensure the policy covers all obligations under the Digital Personal Data Protection Act, 2023 (India), including:
    - Lawful purpose (Sec 4)
    - Informed consent (Sec 6)
    - Notice obligations (Sec 5)
    - Grievance handling (Sec 10)
    - Data accuracy (Sec 8.3)
    - Security safeguards (Sec 8.5)
    - Erasure after purpose is fulfilled or consent withdrawn (Sec 8.7)
    - Special clauses for children (Sec 9) if applicable
    - International transfer declaration (Sec 16)
    
    Write the policy in clear, professional English with practical sections like: Purpose, Scope, Data Types, Lawful Use, Consent, Security, Retention, Rights, Grievance Redressal, Contact.
    
    Return only the policy draft (no disclaimers or titles).
                    """
                    try:
                        draft = call_gpt_text(prompt)
                        st.session_state["full_policy_draft"] = draft
                        st.success("✅ DPDPA-compliant draft generated successfully!")
                    except Exception as e:
                        st.error(f"❌ GPT Error: {e}")
    
        # --- Output Editor ---
        if "full_policy_draft" in st.session_state:
            st.markdown("---")
            st.markdown("### Edit Your Policy")
            edited = st.text_area("Modify the policy text below:", value=st.session_state["full_policy_draft"], height=400, key="full_policy_editor")
    
            col1, col2, col3, col4 = st.columns(4)

            # --- Save in Session ---
            with col1:
                if st.button("💾 Save (Session Only)"):
                    st.session_state["saved_full_policy"] = edited
                    st.success("Draft saved temporarily in session.")
            
            # --- Export to .docx ---
            with col2:
                if st.button("⬇️ Export to Word (.docx)"):
                    doc = Document()
                    doc.add_heading(f"{policy_type} - Generated Policy", level=1)
                    for para in edited.split("\n"):
                        doc.add_paragraph(para)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Word File",
                        data=buffer,
                        file_name=f"{org_name.replace(' ', '_')}_DPDPA_policy.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            # --- Export to .txt ---
            with col3:
                if st.button("⬇️ Export to TXT"):
                    st.download_button(
                        label="📄 Download TXT File",
                        data=edited,
                        file_name=f"{org_name.replace(' ', '_')}_DPDPA_policy.txt",
                        mime="text/plain"
                    )
            
            # --- Save Draft to JSON ---
            with col4:
                if st.button("💾 Save to File (.json)"):
                    draft_data = {
                        "policy": edited,
                        "timestamp": str(datetime.datetime.now()),
                        "org_name": org_name,
                        "policy_type": policy_type
                    }
                    st.download_button(
                        label="📁 Download Draft as JSON",
                        data=json.dumps(draft_data, indent=2),
                        file_name=f"{org_name.replace(' ', '_')}_DPDPA_draft.json",
                        mime="application/json"
                    )

    with tab2:
        with tab2:
            st.markdown("### Generate a Specific Section of Your Policy")
            st.caption("Use this tool to draft a single section aligned with a DPDPA requirement.")
        
            # --- Section Selection ---
            section_map = {
                "Section 4 — Grounds for Processing Personal Data": "4",
                "Section 5 — Notice to Data Principal": "5",
                "Section 6 — Consent & Withdrawal": "6",
                "Section 7 — Legitimate Use Cases": "7",
                "Section 8 — Accuracy, Retention & Security": "8",
                "Section 9 — Processing Children's Data": "9",
                "Section 10 — Grievance Redressal": "10"
            }
        
            st.markdown("**Select DPDPA Section***")
            section_label = st.selectbox("", list(section_map.keys()))
            section_id = section_map[section_label]
        
            # --- Prompt Box ---
            st.markdown("**What part of this section would you like to draft?***")
            st.caption("Example: 'Explain how users can withdraw consent' or 'Write a grievance escalation flow.'")
            custom_instruction = st.text_area("", height=100, key="section_prompt")
        
            # --- Optional Context ---
            st.markdown("**Add Organizational Context (optional)**")
            st.caption("Mention platform name, sector, or audience if needed (e.g., fintech app, KYC users).")
            org_context = st.text_input("", key="section_context")
        
            # --- Generate Button ---
            if st.button("Generate Section"):
                if not custom_instruction.strip():
                    st.warning("Please describe what you want GPT to generate.")
                else:
                    with st.spinner("Generating policy section..."):
                        section_prompt = f"""
        You are a legal assistant drafting a policy section aligned with India's Digital Personal Data Protection Act (DPDPA), 2023.
        
        Draft a clear, compliant, and standalone section for:
        
        **DPDPA Section {section_id} – {section_label.split('—')[-1].strip()}**
        
        Instruction from user: "{custom_instruction.strip()}"
        
        {f'Context: {org_context.strip()}' if org_context.strip() else ''}
        
        Write in plain legal English. Make it usable as-is inside a larger privacy policy.
        Return only the section text. Do not include headings or disclaimers.
                        """
                        try:
                            section_output = call_gpt_text(section_prompt)
                            st.session_state["section_output"] = section_output
                            st.success("✅ Section draft generated successfully!")
                        except Exception as e:
                            st.error(f"❌ GPT Error: {e}")
        
            # --- Output Editor ---
            if "section_output" in st.session_state:
                st.markdown("---")
                st.markdown("### Edit Your Section")
                edited_section = st.text_area("You can make final changes below:", value=st.session_state["section_output"], height=300, key="section_editor")
        
                col1, col2, col3, col4 = st.columns(4)
        
                with col1:
                    if st.button("Save (Session Only)", key="save_section"):
                        st.session_state["saved_section"] = edited_section
                        st.success("Section saved temporarily in session.")
        
                with col2:
                    if st.button("⬇️ Export as Word", key="export_section_docx"):
                        doc = Document()
                        doc.add_heading(section_label, level=1)
                        for para in edited_section.split("\n"):
                            doc.add_paragraph(para)
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="📄 Download Word File",
                            data=buffer,
                            file_name=f"DPDPA_Section_{section_id}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
        
                with col3:
                    if st.button("⬇️ Export as TXT", key="export_section_txt"):
                        st.download_button(
                            label="📄 Download TXT File",
                            data=edited_section,
                            file_name=f"DPDPA_Section_{section_id}.txt",
                            mime="text/plain"
                        )
        
                with col4:
                    if st.button("💾 Save to File (JSON)", key="export_section_json"):
                        section_data = {
                            "section": section_id,
                            "title": section_label,
                            "content": edited_section,
                            "user_instruction": custom_instruction,
                            "org_context": org_context
                        }
                        st.download_button(
                            label="📁 Download JSON Draft",
                            data=json.dumps(section_data, indent=2),
                            file_name=f"DPDPA_Section_{section_id}_Draft.json",
                            mime="application/json"
                        )

    with tab3:
        st.markdown("### Generate Policy by Data Lifecycle Stage")
        st.caption("Use this to generate a specific part of your privacy or retention policy aligned with how data is collected, processed, stored, or shared.")
    
        # --- Lifecycle Options ---
        lifecycle_options = {
            "Data Collection": "Describe what data is collected, from whom, how, and with what consent.",
            "Data Processing": "Explain the purpose and method of processing the data, along with any automation or profiling.",
            "Data Storage": "Detail where data is stored, for how long, and the technical and organizational safeguards.",
            "Data Sharing": "Outline what data is shared, with whom (internal or third party), and under what agreements."
        }
    
        st.markdown("#### Select Lifecycle Stage")
        lifecycle_stage = st.selectbox("", list(lifecycle_options.keys()))
    
        # --- Suggested Template ---
        st.markdown("#### Template Prompt Suggestion")
        st.caption("Modify or use the suggested wording below.")
        default_instruction = lifecycle_options[lifecycle_stage]
        lifecycle_prompt = st.text_area("", value=default_instruction, height=100, key="lifecycle_prompt")
    
        # --- Optional Context ---
        st.markdown("#### Add Organizational Context (optional)")
        st.caption("Include platform name, sector, audience, or known risks (e.g., fintech app handling financial data).")
        lifecycle_context = st.text_input("", key="lifecycle_context")
    
        # --- Generate Section ---
        if st.button("Generate Lifecycle Policy Section"):
            if not lifecycle_prompt.strip():
                st.warning("Please enter or confirm the prompt.")
            else:
                with st.spinner("Generating section..."):
                    lifecycle_prompt_text = f"""
    You are a policy assistant generating a data privacy policy section for a specific lifecycle stage.
    
    **Lifecycle Stage**: {lifecycle_stage}
    **Instruction**: {lifecycle_prompt.strip()}
    {f"Context: {lifecycle_context.strip()}" if lifecycle_context.strip() else ''}
    
    The section should be written in professional, legally sound English and suitable for direct inclusion in a privacy or retention policy. Keep it DPDPA-aligned where applicable.
    Only output the draft content, no explanations or headings.
                    """
                    try:
                        lifecycle_output = call_gpt_text(lifecycle_prompt_text)
                        st.session_state["lifecycle_output"] = lifecycle_output
                        st.success("✅ Section generated successfully!")
                    except Exception as e:
                        st.error(f"❌ GPT Error: {e}")
    
        # --- Output Area ---
        if "lifecycle_output" in st.session_state:
            st.markdown("---")
            st.markdown(f"### Edit Lifecycle Section: {lifecycle_stage}")
            edited_lifecycle = st.text_area("Modify the generated content below:", value=st.session_state["lifecycle_output"], height=300, key="lifecycle_editor")
    
            col1, col2, col3, col4 = st.columns(4)
    
            with col1:
                if st.button("💾 Save (Session Only)", key="save_lifecycle"):
                    st.session_state["saved_lifecycle"] = edited_lifecycle
                    st.success("Saved in session.")
    
            with col2:
                if st.button("⬇️ Export as Word", key="export_lifecycle_docx"):
                    doc = Document()
                    doc.add_heading(f"{lifecycle_stage} Policy", level=1)
                    for para in edited_lifecycle.split("\n"):
                        doc.add_paragraph(para)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Word File",
                        data=buffer,
                        file_name=f"{lifecycle_stage.replace(' ', '_')}_Policy.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
            with col3:
                if st.button("⬇️ Export as TXT", key="export_lifecycle_txt"):
                    st.download_button(
                        label="📄 Download TXT File",
                        data=edited_lifecycle,
                        file_name=f"{lifecycle_stage.replace(' ', '_')}_Policy.txt",
                        mime="text/plain"
                    )
    
            with col4:
                if st.button("💾 Save as JSON", key="export_lifecycle_json"):
                    lifecycle_data = {
                        "stage": lifecycle_stage,
                        "prompt": lifecycle_prompt,
                        "context": lifecycle_context,
                        "content": edited_lifecycle
                    }
                    st.download_button(
                        label="📁 Download JSON Draft",
                        data=json.dumps(lifecycle_data, indent=2),
                        file_name=f"{lifecycle_stage.replace(' ', '_')}_Draft.json",
                        mime="application/json"
                    )


    with tab4:
        st.markdown("### GPT-Assisted Draft Builder")
        st.caption("Describe any policy section or clause you'd like GPT to help you draft. This is not limited to DPDPA or lifecycle templates.")
    
        # --- Prompt Textbox ---
        st.markdown("**Describe what you need help drafting***")
        st.caption("Example: 'Draft a biometric attendance policy for factory workers' or 'Create a privacy notice for children under 13.'")
        free_prompt = st.text_area("", height=120, key="gpt_draft_prompt")
    
        # --- Optional Tags ---
        st.markdown("**Optional Tags**")
        col1, col2, col3 = st.columns(3)
        with col1:
            sector_tag = st.text_input("Sector", help="e.g., Healthcare, Fintech")
        with col2:
            scope_tag = st.text_input("Scope", help="e.g., HR Policy, Vendor Access")
        with col3:
            category_tag = st.text_input("Data Category", help="e.g., Biometric, Financial")
    
        # --- Generate Button ---
        if st.button("Generate with GPT", key="gpt_draft_btn"):
            if not free_prompt.strip():
                st.warning("Please enter a prompt.")
            else:
                with st.spinner("Generating your draft..."):
                    prompt_draft_text = f"""
    You are a policy assistant helping a user draft a professional snippet of policy language.
    
    Instruction: {free_prompt.strip()}
    
    {f"Sector: {sector_tag}" if sector_tag else ''}
    {f"Scope: {scope_tag}" if scope_tag else ''}
    {f"Data Category: {category_tag}" if category_tag else ''}
    
    Write in clear, professional policy language. Avoid filler text, disclaimers, or general advice. Return only the content of the policy.
                    """
                    try:
                        gpt_draft_output = call_gpt_text(prompt_draft_text)
                        st.session_state["gpt_draft_output"] = gpt_draft_output
                        st.success("✅ Draft generated!")
                    except Exception as e:
                        st.error(f"❌ GPT Error: {e}")
    
        # --- Editable Output Area ---
        if "gpt_draft_output" in st.session_state:
            st.markdown("---")
            st.markdown("### Edit Your Draft")
            edited_gpt_draft = st.text_area("Make final changes below:", value=st.session_state["gpt_draft_output"], height=300, key="gpt_draft_editor")
    
            col1, col2, col3, col4 = st.columns(4)
    
            with col1:
                if st.button("💾 Save (Session Only)", key="save_gpt_draft"):
                    st.session_state["saved_gpt_draft"] = edited_gpt_draft
                    st.success("Saved in session.")
    
            with col2:
                if st.button("⬇️ Export as Word", key="export_gpt_draft_docx"):
                    doc = Document()
                    doc.add_heading("Custom Policy Draft", level=1)
                    for para in edited_gpt_draft.split("\n"):
                        doc.add_paragraph(para)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Word File",
                        data=buffer,
                        file_name=f"Custom_Policy_Draft.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
            with col3:
                if st.button("⬇️ Export as TXT", key="export_gpt_draft_txt"):
                    st.download_button(
                        label="📄 Download TXT File",
                        data=edited_gpt_draft,
                        file_name=f"Custom_Policy_Draft.txt",
                        mime="text/plain"
                    )
    
            with col4:
                if st.button("💾 Save as JSON", key="export_gpt_draft_json"):
                    gpt_draft_data = {
                        "prompt": free_prompt,
                        "sector": sector_tag,
                        "scope": scope_tag,
                        "category": category_tag,
                        "content": edited_gpt_draft
                    }
                    st.download_button(
                        label="📁 Download JSON Draft",
                        data=json.dumps(gpt_draft_data, indent=2),
                        file_name="Custom_Policy_Draft.json",
                        mime="application/json"
                    )


    with tab5:
        st.markdown("### View & Manage Saved Drafts")
        st.caption("This page lets you view, edit, rename, export, or delete any draft stored in your session or uploaded from a file.")
    
        # --- Session State: Collect all saved drafts ---
        saved_drafts = {}
        for key in st.session_state:
            if key.startswith("saved_"):
                label = key.replace("saved_", "").replace("_", " ").title()
                saved_drafts[label] = st.session_state[key]
    
        # --- Upload JSON Draft ---
        st.markdown("#### Upload a JSON Draft (from export)")
        uploaded = st.file_uploader("Upload JSON file", type="json")
        if uploaded:
            try:
                loaded = json.load(uploaded)
                label = f"Uploaded: {uploaded.name}"
                saved_drafts[label] = loaded["content"]
                st.success(f"✅ Loaded draft: {label}")
            except Exception as e:
                st.error(f"❌ Error loading draft: {e}")
    
        if not saved_drafts:
            st.info("No drafts found in this session. Try generating and saving one in Tabs 1–4.")
        else:
            draft_names = list(saved_drafts.keys())
            selected = st.selectbox("📋 Select a draft to view/edit", draft_names)
            current_draft = saved_drafts[selected]
    
            st.markdown(f"#### Editing Draft: `{selected}`")
            edited_draft = st.text_area("You can update the draft content below:", value=current_draft, height=300, key="edit_draft")
    
            col1, col2, col3, col4 = st.columns(4)
    
            # --- Rename ---
            with col1:
                new_name = st.text_input("Rename this draft as", value=selected, key="rename_field")
                if st.button("🔁 Rename Draft"):
                    st.session_state[f"saved_{new_name.replace(' ', '_').lower()}"] = edited_draft
                    if f"saved_{selected.replace(' ', '_').lower()}" in st.session_state:
                        del st.session_state[f"saved_{selected.replace(' ', '_').lower()}"]
                    st.success("✅ Draft renamed successfully. Refresh the dropdown to see updated name.")
    
            # --- Export .docx ---
            with col2:
                if st.button("⬇️ Export as Word (.docx)", key="export_word_saved"):
                    doc = Document()
                    doc.add_heading(f"{selected} Draft", level=1)
                    for para in edited_draft.split("\n"):
                        doc.add_paragraph(para)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Word File",
                        data=buffer,
                        file_name=f"{selected.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
            # --- Export .txt ---
            with col3:
                if st.button("⬇️ Export as TXT", key="export_txt_saved"):
                    st.download_button(
                        label="📄 Download TXT File",
                        data=edited_draft,
                        file_name=f"{selected.replace(' ', '_')}.txt",
                        mime="text/plain"
                    )
    
            # --- Export .json ---
            with col4:
                if st.button("💾 Download JSON", key="export_json_saved"):
                    st.download_button(
                        label="📁 Download JSON File",
                        data=json.dumps({"name": selected, "content": edited_draft}, indent=2),
                        file_name=f"{selected.replace(' ', '_')}.json",
                        mime="application/json"
                    )
    
            # --- Delete from session ---
            if st.button("🗑️ Delete Draft", key="delete_draft_btn"):
                del_key = f"saved_{selected.replace(' ', '_').lower()}"
                if del_key in st.session_state:
                    del st.session_state[del_key]
                    st.success(f"Draft `{selected}` deleted. Refresh to confirm.")


# --- Policy Compliance Checker ---
elif menu == "Policy Compliance Checker":

    #st.title("Match Policy to DPDPA")
    st.markdown("<h1 style='font-size:38px; font-weight:800;'>Match Policy to DPDPA</h1>", unsafe_allow_html=True)
    
    #st.header("1. Upload Your Policy Document")
    st.markdown("<h3 style='font-size:24px; font-weight:700;'>1. Upload Your Policy Document</h3>", unsafe_allow_html=True)

    upload_option = st.radio("Choose input method:", ["Paste text", "Upload PDF"])
    if upload_option == "Paste text":
        policy_text = st.text_area("Paste your Privacy Policy text:", height=300)
    elif upload_option == "Upload PDF":
        uploaded_pdf = st.file_uploader("Upload PDF file", type="pdf", label_visibility="collapsed")

        if uploaded_pdf:
            # 👇 Custom visible filename
            st.markdown(f"""
            <div style="padding: 6px 12px; background-color:#f1f1f1; display:inline-block;
                        border-radius:6px; font-weight:600; color:#000; font-family: Arial, sans-serif;">
            📄 Uploaded file: {uploaded_pdf.name}
            </div>
            """, unsafe_allow_html=True)

            policy_text = extract_text_from_pdf(uploaded_pdf)
        else:
            policy_text = ""

    #st.header("4. Industry Context (Optional)")
    st.markdown("<h3 style='font-size:24px; font-weight:700;'>2. Industry Context (Optional)</h3>", unsafe_allow_html=True)
    industry = st.selectbox("", ["General", "Automotive", "Healthcare", "Fintech", "Other"])
    if industry == "Other":
        custom_industry = st.text_input("Specify your industry")
    else:
        custom_industry = None

    st.markdown("<h3 style='font-size:24px; font-weight:700;'>3. Choose DPDPA Section</h3>", unsafe_allow_html=True)
    # section_options = list(dpdpa_checklists.keys()) + ["All Sections"]
    section_options = [f"{sid} — {dpdpa_checklists[sid]['title']}" for sid in dpdpa_checklists] + ["All Sections"]
    section_id = st.selectbox("", options=section_options)

    st.markdown("<h3 style='font-size:24px; font-weight:700;'>4. Run Compliance Check</h3>", unsafe_allow_html=True)
    if st.button("Run Compliance Check"):
        if policy_text:
            result = []
            with st.spinner("Running GPT-based compliance evaluation..."):
                if section_id == "All Sections":
                    all_results = []  # 🔁 collect each section's result

                    for sid in dpdpa_checklists:
                        st.markdown(f"## ✅ Processing Section {sid} — {dpdpa_checklists[sid]['title']}")
                        checklist = dpdpa_checklists[sid]["items"]
                        result = analyze_policy_section(sid, checklist, policy_text)
                        all_results.append(result)
            
                        with st.expander(f"Section {result['Section']} — {result['Title']}", expanded=True):
                            level_color = {
                                "Fully Compliant": "#198754",
                                "Partially Compliant": "#FFC107",
                                "Non-Compliant": "#DC3545"
                            }
                            match_level = result["Match Level"]
                            color = level_color.get(match_level, "#6C757D")
            
                            st.markdown(f"""
                            <div style="margin-bottom: 1rem;">
                              <b>Compliance Score:</b>
                              <span style="background-color:#0d6efd; color:white; padding:4px 10px; border-radius:5px; font-size:0.85rem;">
                                {result["Compliance Score"]}
                              </span><br>
                              <b>Match Level:</b>
                              <span style="background-color:{color}; color:black; padding:4px 10px; border-radius:5px; font-size:0.85rem;">
                                {match_level}
                              </span>
                            </div>
                            """, unsafe_allow_html=True)
            
                            st.markdown("### 📋 Checklist Items Matched:")
                            for item in result["Checklist Items Matched"]:
                                st.markdown(f"- {item}")
            
                            st.markdown("### 🔍 Matched Details:")
                            for item in result["Matched Details"]:
                                status = item.get("Status", "Missing")
                                badge_color = {
                                    "Explicitly Mentioned": "#198754",
                                    "Partially Mentioned": "#FFC107",
                                    "Missing": "#DC3545"
                                }.get(status, "#6c757d")
            
                                st.markdown(f"""
                                **{item['Checklist Item ID']} — {item['Checklist Text']}**  
                                <span style="color:white;background-color:{badge_color};padding:3px 10px;border-radius:6px;font-size:13px;">{status}</span>  
                                <br><small>📝 {item.get("Justification", "No justification")}</small>
                                """, unsafe_allow_html=True)
            
                            st.markdown("### ✏️ Suggested Rewrite:")
                            st.info(result["Suggested Rewrite"])
            
                            st.markdown("### 🧾 Simplified Legal Meaning:")
                            st.success(result["Simplified Legal Meaning"])
            
                    # ✅ Combined Export Section
                    st.markdown("## 📥 Export Combined Results")
            
                    # --- JSON Export ---
                    combined_json = json.dumps(all_results, indent=2)
                    combined_json_bytes = io.BytesIO(combined_json.encode("utf-8"))
                    st.download_button(
                        label="📥 Download Combined JSON",
                        data=combined_json_bytes,
                        file_name="DPDPA_All_Sections_Combined.json",
                        mime="application/json"
                    )
            
                    # --- CSV Export ---
                    combined_rows = []
                    for result in all_results:
                        for item in result["Matched Details"]:
                            combined_rows.append({
                                "Section": result["Section"],
                                "Checklist Item ID": item["Checklist Item ID"],
                                "Checklist Text": item["Checklist Text"],
                                "Status": item["Status"],
                                "Justification": item["Justification"],
                                "Match Level": result["Match Level"],
                                "Score": result["Compliance Score"]
                            })
            
                    combined_csv = pd.DataFrame(combined_rows)
                    combined_csv_bytes = io.BytesIO()
                    combined_csv.to_csv(combined_csv_bytes, index=False)
                    combined_csv_bytes.seek(0)
            
                    st.download_button(
                        label="📥 Download Combined CSV",
                        data=combined_csv_bytes,
                        file_name="DPDPA_All_Sections_Evaluation.csv",
                        mime="text/csv"
                    )
                else:
                    section_num = section_id.split(" — ")[0] if " — " in section_id else section_id
                    checklist = dpdpa_checklists[section_num]['items']

                    result = analyze_policy_section(section_num, checklist, policy_text)
                    st.markdown(f"""
                    <div style='font-size:20px; font-weight:700; margin-top:25px; margin-bottom:-10px;'>
                    📘 Section {result['Section']} — {result['Title']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    with st.expander("", expanded=True):

                        # Set color for Match Level badge
                        level_color = {
                            "Fully Compliant": "#198754",     # green
                            "Partially Compliant": "#FFC107", # yellow
                            "Non-Compliant": "#DC3545"        # red
                        }
                        match_level = result["Match Level"]
                        color = level_color.get(match_level, "#6C757D")  # fallback grey
                        
                        st.markdown(f"""
                        <div style="margin-bottom: 1rem;">
                          <b>Compliance Score:</b>
                          <span style="background-color:#0d6efd; color:white; padding:4px 10px; border-radius:5px; font-size:0.85rem;">
                            {result["Compliance Score"]}
                          </span><br>
                          <b>Match Level:</b>
                          <span style="background-color:{color}; color:black; padding:4px 10px; border-radius:5px; font-size:0.85rem;">
                            {match_level}
                          </span>
                        </div>
                        """, unsafe_allow_html=True)

                    
                        st.markdown("### 📋 Checklist Items Matched:")
                        for i, item in enumerate(result["Checklist Items Matched"]):
                            st.markdown(f"- {item}")
            
                        st.markdown("### 🔍 Matched Details:")
                        for item in result["Matched Details"]:
                            status = item.get("Status", "Missing")
                            color = {
                                "Explicitly Mentioned": "#198754",
                                "Partially Mentioned": "#FFC107",
                                "Missing": "#DC3545"
                            }.get(status, "#6c757d")
                        
                            item_id = item.get("Checklist Item ID", "❓")
                            item_text = item.get("Checklist Text", "❓")
                            justification = item.get("Justification", "No justification found.")
                        
                            st.markdown(f"""
                        **{item_id} — {item_text}**  
                        <span style="color:white;background-color:{color};padding:3px 10px;border-radius:6px;font-size:13px;">{status}</span>  
                        <br><small>📝 {justification}</small>
                        """, unsafe_allow_html=True)
                    
                        st.markdown("### ✏️ Suggested Rewrite:")
                        st.info(result["Suggested Rewrite"])
                    
                        st.markdown("### 🧾 Simplified Legal Meaning:")
                        st.success(result["Simplified Legal Meaning"])

                        # --- JSON Export ---
                        json_str = json.dumps(result, indent=2)
                        json_bytes = io.BytesIO(json_str.encode("utf-8"))
                        st.download_button(
                            label="📥 Download JSON Report",
                            data=json_bytes,
                            file_name=f"DPDPA_Section_{result['Section']}.json",
                            mime="application/json"
                        )
                        
                        # --- CSV Export ---
                        csv_df = pd.DataFrame(result["Matched Details"])
                        csv_bytes = io.BytesIO()
                        csv_df.to_csv(csv_bytes, index=False)
                        csv_bytes.seek(0)
                        st.download_button(
                            label="📥 Download Checklist Evaluation CSV",
                            data=csv_bytes,
                            file_name=f"DPDPA_Section_{result['Section']}.csv",
                            mime="text/csv"
                        )
